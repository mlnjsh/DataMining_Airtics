"""Convert EDA_Report.md to EDA_Report.docx (headings, bullets, tables, images, bold)."""
import re, os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

SRC, OUT = "EDA_Report.md", "EDA_Report.docx"
doc = Document()
for s in doc.sections:
    s.left_margin = s.right_margin = Inches(0.9)
    s.top_margin = s.bottom_margin = Inches(0.8)
doc.styles["Normal"].font.name = "Calibri"
doc.styles["Normal"].font.size = Pt(11)


def add_runs(par, text):
    """Handle **bold** and `code` inline."""
    for part in re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", text):
        if not part:
            continue
        if part.startswith("**"):
            r = par.add_run(part[2:-2]); r.bold = True
        elif part.startswith("`"):
            r = par.add_run(part[1:-1]); r.font.name = "Consolas"; r.font.size = Pt(10)
        else:
            par.add_run(part)


def add_table(rows):
    header, body = rows[0], rows[2:]          # rows[1] is the |---| line
    t = doc.add_table(rows=1 + len(body), cols=len(header))
    t.style = "Light Grid Accent 1"
    for j, h in enumerate(header):
        c = t.rows[0].cells[j]; c.text = ""
        r = c.paragraphs[0].add_run(h); r.bold = True; r.font.size = Pt(10)
    for i, row in enumerate(body, start=1):
        for j in range(len(header)):
            val = row[j] if j < len(row) else ""
            c = t.rows[i].cells[j]; c.text = ""
            add_runs(c.paragraphs[0], val)
            for r in c.paragraphs[0].runs: r.font.size = Pt(10)
    doc.add_paragraph()


lines = open(SRC, encoding="utf-8").read().split("\n")
i = 0
while i < len(lines):
    ln = lines[i]
    if ln.startswith("|"):
        rows = []
        while i < len(lines) and lines[i].startswith("|"):
            rows.append([c.strip() for c in lines[i].strip().strip("|").split("|")]); i += 1
        add_table(rows); continue
    m = re.match(r"^(#+)\s+(.*)", ln)
    if m:
        level = len(m.group(1))
        doc.add_heading(m.group(2), level=min(level, 3))
    elif ln.startswith("!["):
        path = re.search(r"\((.*?)\)", ln).group(1)
        if os.path.exists(path):
            doc.add_picture(path, width=Inches(6.2))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif ln.startswith("- "):
        add_runs(doc.add_paragraph(style="List Bullet"), ln[2:])
    elif re.match(r"^\d+\.\s", ln):
        add_runs(doc.add_paragraph(style="List Number"), re.sub(r"^\d+\.\s", "", ln))
    elif ln.strip() == "---":
        doc.add_page_break()
    elif ln.strip():
        add_runs(doc.add_paragraph(), ln)
    i += 1

doc.save(OUT)
print("saved", OUT)
