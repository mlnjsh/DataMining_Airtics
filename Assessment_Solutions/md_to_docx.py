"""
Convert a Markdown report to a Word document.
Handles headings, paragraphs, bullet and numbered lists, tables,
images, fenced code blocks, and inline **bold** / `code`.

Usage:  python md_to_docx.py <input.md> [output.docx]
Images are resolved relative to the Markdown file.
"""
import re, os, sys
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

src = sys.argv[1]
out = sys.argv[2] if len(sys.argv) > 2 else os.path.splitext(src)[0] + ".docx"
base = os.path.dirname(os.path.abspath(src))

doc = Document()
for s in doc.sections:
    s.left_margin = s.right_margin = Inches(0.9)
    s.top_margin = s.bottom_margin = Inches(0.8)
doc.styles["Normal"].font.name = "Calibri"
doc.styles["Normal"].font.size = Pt(11)


def shade(cell, hex_fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), hex_fill)
    tcPr.append(shd)


def add_runs(par, text):
    for part in re.split(r"(\*\*[^*]+\*\*|`[^`]+`|\*[^*]+\*)", text):
        if not part:
            continue
        if part.startswith("**"):
            r = par.add_run(part[2:-2]); r.bold = True
        elif part.startswith("`"):
            r = par.add_run(part[1:-1]); r.font.name = "Consolas"; r.font.size = Pt(10)
        elif part.startswith("*") and len(part) > 2:
            r = par.add_run(part[1:-1]); r.italic = True
        else:
            par.add_run(part)


def add_table(rows):
    header = rows[0]
    body = [r for r in rows[1:] if not re.match(r"^:?-{2,}", r[0].strip() or "-")]
    body = [r for r in rows[1:] if not all(re.match(r"^:?-+:?$", c.strip()) or c.strip() == "" for c in r)]
    ncol = len(header)
    t = doc.add_table(rows=1 + len(body), cols=ncol)
    t.style = "Table Grid"
    for j, h in enumerate(header):
        c = t.rows[0].cells[j]; c.text = ""
        r = c.paragraphs[0].add_run(h); r.bold = True; r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF); shade(c, "2E5A88")
    for i, row in enumerate(body, start=1):
        for j in range(ncol):
            val = row[j] if j < len(row) else ""
            c = t.rows[i].cells[j]; c.text = ""
            add_runs(c.paragraphs[0], val)
            for r in c.paragraphs[0].runs: r.font.size = Pt(10)
            if i % 2 == 0: shade(c, "F2F5F9")
    doc.add_paragraph()


def add_code_block(code_lines):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), "F2F5F9")
    pPr.append(shd)
    for k, line in enumerate(code_lines):
        r = p.add_run(line)
        r.font.name = "Consolas"; r.font.size = Pt(9.5)
        if k < len(code_lines) - 1:
            r.add_break()


lines = open(src, encoding="utf-8").read().split("\n")
i = 0
while i < len(lines):
    ln = lines[i]

    if ln.startswith("```"):
        i += 1; block = []
        while i < len(lines) and not lines[i].startswith("```"):
            block.append(lines[i]); i += 1
        add_code_block(block); i += 1; continue

    if ln.startswith("|"):
        rows = []
        while i < len(lines) and lines[i].startswith("|"):
            rows.append([c.strip() for c in lines[i].strip().strip("|").split("|")]); i += 1
        add_table(rows); continue

    m = re.match(r"^(#+)\s+(.*)", ln)
    if m:
        doc.add_heading(m.group(2).strip(), level=min(len(m.group(1)), 3))
    elif ln.startswith("!["):
        path = os.path.join(base, re.search(r"\((.*?)\)", ln).group(1))
        if os.path.exists(path):
            doc.add_picture(path, width=Inches(6.2))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif re.match(r"^\s*[-*]\s", ln):
        indent = len(ln) - len(ln.lstrip())
        p = doc.add_paragraph(style="List Bullet 2" if indent >= 2 else "List Bullet")
        add_runs(p, re.sub(r"^\s*[-*]\s", "", ln))
    elif re.match(r"^\d+\.\s", ln):
        add_runs(doc.add_paragraph(style="List Number"), re.sub(r"^\d+\.\s", "", ln))
    elif ln.strip() == "---":
        doc.add_page_break()
    elif ln.strip():
        add_runs(doc.add_paragraph(), ln)
    i += 1

doc.save(out)
print("saved", out)
